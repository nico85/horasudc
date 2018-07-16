# -*- coding: utf-8 -*-
from django.http import HttpResponse
from django.http import HttpResponseRedirect
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.staticfiles import finders
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger

from horas.forms import PersonaForm, PersonaHorasForm, DocenteHorasForm

from horas.models import *

from docx import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.shared import Cm

import csv
import math

from datetime import date

# Create your views here.

gbl_cons_adm = [None]
gbl_cons_doc = [None]

@login_required()
def inicio(request):
    version = Version.objects.last()
    cambios = Cambio.objects.filter(version_id=version.id)
    return render(request, 'inicio.html', {
        'version': version,
        'cambios': cambios
    })

@login_required()
def personasList(request):

    #grupo = request.user.groups.get()
    grupo = 'superusuario'
    todas_las_personas = Persona.objects.all().order_by('apellidos')
    #docentes = DocenteHoras.objects.all()
    #administrativos = PersonaHoras.objects.all()

    personas = todas_las_personas
    #cantPersonas = todas_las_personas.count()

    if request.method == 'GET':

        if len(todas_las_personas) > 0:

            '''paginator = Paginator(todas_las_personas, 20)  # Show 20 contacts per page
    
            page = request.GET.get('page', 1)
    
            try:
                personas = paginator.page(page)
            except PageNotAnInteger:
                # If page is not an integer, deliver first page.
                personas = paginator.page(1)
            except EmptyPage:
                # If page is out of range (e.g. 9999), deliver last page of results.
                personas = paginator.page(paginator.num_pages)
            '''
            return render(request, 'personasList.html', {
                #'page': page,
                #'paginator': paginator,
                'personas': personas,
                'grupo': grupo,
                #'docentes': docentes,
                #'administrativos': administrativos,
                #'cantPersonas': cantPersonas,
                'estadoSel': -1
            })
        else:
            return render(request, 'personasList.html', {
                'personas': personas,
                'grupo': grupo,
                #'docentes': docentes,
                #'administrativos': administrativos,
                #'cantPersonas': cantPersonas,
            })
    else:
        estadoSel = -1
        if 'estado' in request.POST and int(request.POST['estado']) > -1:
            personas = personas.filter(activo=request.POST['estado'])
            estadoSel = int(request.POST['estado'])
            cantPersonas = len(personas)
        return render(request, 'personasList.html', {
            # 'page': page,
            # 'paginator': paginator,
            'personas': personas,
            'grupo': grupo,
            #'docentes': docentes,
            #'administrativos': administrativos,
            #'cantPersonas': cantPersonas,
            'estadoSel': estadoSel,
        })

@login_required()
def personasNew(request):
    lista_sexo = ['Femenino', 'Masculino']
    if request.method == 'POST':
        form = PersonaForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('/personas/')
        else:
            return render(request, 'personasNew.html', {
                'form': form,
                'sexos': lista_sexo,
            })
    else:
        form = PersonaForm()
        return render(request, 'personasNew.html', {
            'form': form,
            'sexos': lista_sexo,
        })


@login_required()
def personasEdit(request, pid):
    persona = get_object_or_404(Persona, id=pid)
    lista_sexo = ['Femenino', 'Masculino']
    if request.method == 'POST':
        form = PersonaForm(request.POST, instance=persona)
        if form.is_valid():
            form.save()
            return redirect('/personas/')
        else:
            return render(request, 'personasEdit.html', {
                'persona': persona,
                'form': form,
                'sexos': lista_sexo,
            })
    else:
        form = PersonaForm(instance=persona)
        return render(request, 'personasEdit.html', {
            'persona': persona,
            'form': form,
            'sexos': lista_sexo,
        })

@login_required()
def personasVer(request, pid):
    persona = get_object_or_404(Persona, id=pid)
    return render(request, 'personasVer.html', {
        'persona': persona
    })

@login_required()
def personasCertPrestServ(request, pid):
    persona = get_object_or_404(Persona, id=pid)
    docHoras = DocenteHoras.objects.filter(persona_id=pid).order_by('fecha_inicio')
    persHoras = PersonaHoras.objects.filter(persona_id=pid).order_by('fecha_inicio')

    return render(request, 'personasCertifPrestServ.html', {
        'persona': persona,
        'persHoras': persHoras,
        'docHoras': docHoras,
    })

@login_required()
def consadm(request):
    # obtener el Arreglo de años (actual al 2009)
    hoy = date.today()
    anio_actual = hoy.year
    nro_fil = (anio_actual - 2008)
    lista_anios = [None] * nro_fil
    i = 0
    for n in range(nro_fil):
        lista_anios[n] = str(anio_actual - i)
        i += 1
    #obtengo el total de las asignaciones de horas de administrativos
    adms = PersonaHoras.objects.all().order_by('-fecha_fin')
    total_adms = adms.count()
    total_adms_baja = adms.filter(baja=True).count()
    #obtengo todas las dependencias
    dependencias = Dependencia.objects.all()
    anioSel = '0'
    dependenciaSel = '0'
    fechaDesdeSel = ''
    fechaHastaSel = ''
    fechaDesdeVigSel = ''
    fechaHastaVigSel = ''
    resnro = 0
    isGet = False

    if request.method == 'GET':
        isGet = True
        return render(request, 'consultaadministrativo.html', {
            'administrativos': adms,
            'dependencias': dependencias,
            'anios': lista_anios,
            'total_adms': total_adms,
            'total_adms_baja': total_adms_baja,
            'resnro': resnro,
            'isGet': isGet
        })        
    else:
        adms = PersonaHoras.objects.all()
        admsBaja = adms.filter(baja=True)
        adms = adms.filter(baja=False)
        anio = 0
        dependencia = 0
        #si viene filtro por año
        if 'anio' in request.POST and int(request.POST['anio']) > 0:
            adms = adms.filter(resolucion_anio=request.POST['anio'])
            admsBaja = admsBaja.filter(resolucion_anio=request.POST['anio'])
            anioSel = request.POST['anio']
        #si viene filtro por numero de resolucion
        if 'res_nro' in request.POST and request.POST['res_nro'] != "":
            adms = adms.filter(resolucion_numero=request.POST['res_nro'])
            admsBaja = admsBaja.filter(resolucion_numero=request.POST['res_nro'])
            resnro = request.POST['res_nro']
        # si viene filtro por dependencia
        if 'dependencia' in request.POST and int(request.POST['dependencia']) > 0:
            adms = adms.filter(dependencia_id=request.POST['dependencia'])
            admsBaja = admsBaja.filter(dependencia_id=request.POST['dependencia'])
            dependenciaSel = int(request.POST['dependencia'])
        # si viene filtro por fecha desde
        #if 'fecha_desde' in request.POST and request.POST['fecha_desde'] != "":
        #    fechaDesdeSel = request.POST['fecha_desde']
        #    if 'fecha_hasta' in request.POST and request.POST['fecha_hasta'] != "":
        #        # ejemplo: Order.objects.filter(drop_off__gte=start_date, pick_up__lte=end_date)
        #        #__gte: greater-than | __lte: lees-than
        #        adms = adms.filter(fecha_inicio__gte=request.POST['fecha_desde'], fecha_fin__lte=request.POST['fecha_hasta'])
        #        admsBaja = admsBaja.filter(fecha_inicio__gte=request.POST['fecha_desde'],
        #                           fecha_baja__lte=request.POST['fecha_hasta'])
        #        fechaHastaSel = request.POST['fecha_hasta']
        #    else:
        #        adms = adms.filter(fecha_inicio__gte=request.POST['fecha_desde'])
        #        admsBaja = admsBaja.filter(fecha_inicio__gte=request.POST['fecha_desde'])
        #else:
        #    if 'fecha_hasta' in request.POST and request.POST['fecha_hasta'] != "":
        #        adms = adms.filter(fecha_fin__lte=request.POST['fecha_hasta'])
        #        admsBaja = admsBaja.filter(fecha_fin__lte=request.POST['fecha_hasta'])
        #        fechaHastaSel = request.POST['fecha_hasta']
        
        #si viene fecha_desde_vigente?
        if 'fecha_desde_v' in request.POST and request.POST['fecha_desde_v'] != "":
			#guardo la variable para mostrar en la vista luego
            fechaDesdeVigSel = request.POST['fecha_desde_v']
            #me fijo si viene fecha_hasta_vigente
            if 'fecha_hasta_v' in request.POST and request.POST['fecha_hasta_v'] != "":
                adms = adms.filter(fecha_inicio__lte=request.POST['fecha_desde_v'],
                                   fecha_fin__gte=request.POST['fecha_hasta_v'])
                admsBaja = admsBaja.filter(fecha_inicio__lte=request.POST['fecha_desde_v'],
                                   fecha_baja__gte=request.POST['fecha_hasta_v'])
                fechaHastaVigSel = request.POST['fecha_hasta_v']
            else:
                adms = adms.filter(fecha_inicio__lte=request.POST['fecha_desde_v'])
                admsBaja = admsBaja.filter(fecha_inicio__lte=request.POST['fecha_desde_v'])
        else:
            if 'fecha_hasta_v' in request.POST and request.POST['fecha_hasta_v'] != "":
                adms = adms.filter(fecha_fin__gte=request.POST['fecha_hasta_v'])
                admsBaja = admsBaja.filter(fecha_fin__gte=request.POST['fecha_hasta_v'])
                fechaHastaVigSel = request.POST['fecha_hasta_v']

        total_adms_baja = admsBaja.count()
        administrativos = adms | admsBaja
        administrativos = administrativos.distinct().order_by('-fecha_fin')

        global gbl_cons_adm
        gbl_cons_adm = administrativos

        return render(request, 'consultaadministrativo.html', {
            'administrativos': administrativos,
            'dependencias': dependencias,
            'anios': lista_anios,
            'dependencia': dependencia,
            'anio': anio,
            'resnro': resnro,
            'total_adms': total_adms,
            'total_adms_baja': total_adms_baja,
            'anioSel': anioSel,
            'isGet': isGet,
            'dependenciaSel': dependenciaSel,
            'fechaDesdeSel': fechaDesdeSel,
            'fechaHastaSel': fechaHastaSel,
            'fechaDesdeVigSel': fechaDesdeVigSel,
            'fechaHastaVigSel': fechaHastaVigSel
        })


@login_required()
def consdoc(request):
    # obtener el Arreglo de años (actual al 2009)
    hoy = date.today()
    anio_actual = hoy.year
    nro_fil = (anio_actual - 2008)
    lista_anios = [None] * nro_fil
    i = 0
    for n in range(nro_fil):
        lista_anios[n] = str(anio_actual - i)
        i += 1
    #obtengo todas las sedes
    sedes = Sede.objects.all()
    #obtengo todas las carerras
    carreras = Carrera.objects.all()
    # obtengo el total de las asignaciones de horas de Docentes
    docs = DocenteHoras.objects.all().order_by('-fecha_fin')
    #obtengo los totales
    total_docs = docs.count()
    total_docs_baja = docs.filter(baja=True).count()
    sedeSel = 0
    carreraSel = 0
    anioSel = '0'
    fechaDesdeSel = ''
    fechaHastaSel = ''
    fechaDesdeVigSel = ''
    fechaHastaVigSel = ''
    resnro = 0
    isGet = False

    nro_fil = len(docs)
    lista_hs = [None] * nro_fil
    for n in range(nro_fil):
        lista_hs[n] = {'id_doc': 0, 'tot_hs': 0}
        i = 0
    for dh in docs:
        tot_pocentaje = (dh.porcentaje_aplicado * dh.materia.hs_semanales) / 100
        tot = math.ceil(dh.materia.hs_semanales + dh.hs_institucionales + tot_pocentaje)
        lista_hs[i]['id_doc'] = dh.id
        lista_hs[i]['tot_hs'] = int(tot)
        i += 1

    if request.method == 'GET':
        isGet = True
        return render(request, 'consultadocente.html', {
            'docentes': docs,
            'sedes': sedes,
            'carreras': carreras,
            'anios': lista_anios,
            'resnro': resnro,
            'total_docs': total_docs,
            'total_docs_baja': total_docs_baja,
            'isGet': isGet,
            'lista_hs': lista_hs,
        })
    else:
        docs = DocenteHoras.objects.all()
        docs_baja = docs.filter(baja=True)
        docs = docs.filter(baja=False)
        anio = 0
        sede = 0
        carrera = 0
        #si viene para dfiltrar por sede
        if 'sede' in request.POST and int(request.POST['sede']) > 0:
            docs = docs.filter(sede=request.POST['sede'])
            docs_baja = docs_baja.filter(sede=request.POST['sede'])
            sedeSel = int(request.POST['sede'])
        # si viene para dfiltrar por carrera
        if 'carrera' in request.POST and int(request.POST['carrera']) > 0:
            docs = docs.filter(materia__plan__carrera_id=request.POST['carrera'])
            docs_baja = docs_baja.filter(materia__plan__carrera_id=request.POST['carrera'])
            carreraSel = int(request.POST['carrera'])
        # si viene para dfiltrar por año
        if 'anio' in request.POST and int(request.POST['anio']) > 0:
            docs = docs.filter(resolucion_anio=request.POST['anio'])
            docs_baja = docs_baja.filter(resolucion_anio=request.POST['anio'])
            anioSel = request.POST['anio']
        # si viene para dfiltrar por numero de resolucion
        if 'res_nro' in request.POST and request.POST['res_nro'] != "":
            docs = docs.filter(resolucion_numero=request.POST['res_nro'])
            docs_baja = docs_baja.filter(resolucion_numero=request.POST['res_nro'])
            resnro = request.POST['res_nro']
        # si viene para dfiltrar por numero Fechas (desde y Hasta
        #if 'fecha_desde' in request.POST and request.POST['fecha_desde'] != "":
        #    fechaDesdeSel = request.POST['fecha_desde']
        #    if 'fecha_hasta' in request.POST and request.POST['fecha_hasta'] != "":
        #        # ejemplo: Order.objects.filter(drop_off__gte=start_date, pick_up__lte=end_date)
        #        docs = docs.filter(fecha_inicio__gte=request.POST['fecha_desde'], fecha_fin__lte=request.POST['fecha_hasta'])
        #        fechaHastaSel = request.POST['fecha_hasta']
        #    else:
        #        docs = docs.filter(fecha_inicio__gte=request.POST['fecha_desde'])
        #else:
        #    if 'fecha_hasta' in request.POST and request.POST['fecha_hasta'] != "":
        #        docs = docs.filter(fecha_fin__lte=request.POST['fecha_hasta'])
        #        fechaHastaSel = request.POST['fecha_hasta']

        # si viene fecha_desde_vigente?
        if 'fecha_desde_v' in request.POST and request.POST['fecha_desde_v'] != "":
            # guardo la variable para mostrar en la vista luego
            fechaDesdeVigSel = request.POST['fecha_desde_v']
            # me fijo si viene fecha_hasta_vigente
            if 'fecha_hasta_v' in request.POST and request.POST['fecha_hasta_v'] != "":
                docs = docs.filter(fecha_inicio__lte=request.POST['fecha_desde_v'],
                                   fecha_fin__gte=request.POST['fecha_hasta_v'])
                docs_baja = docs_baja.filter(fecha_inicio__lte=request.POST['fecha_desde_v'],
                                           fecha_baja__gte=request.POST['fecha_hasta_v'])
                fechaHastaVigSel = request.POST['fecha_hasta_v']
            else:
                docs = docs.filter(fecha_inicio__lte=request.POST['fecha_desde_v'])
                docs_baja = docs_baja.filter(fecha_inicio__lte=request.POST['fecha_desde_v'])
        else:
            if 'fecha_hasta_v' in request.POST and request.POST['fecha_hasta_v'] != "":
                docs = docs.filter(fecha_fin__gte=request.POST['fecha_hasta_v'])
                docs_baja = docs_baja.filter(fecha_fin__gte=request.POST['fecha_hasta_v'])
                fechaHastaVigSel = request.POST['fecha_hasta_v']

        total_docs_baja = docs_baja.count()
        docentes = docs | docs_baja
        docentes = docentes.distinct()

        global gbl_cons_doc
        gbl_cons_doc = docentes

        return render(request, 'consultadocente.html', {
            'docentes': docentes,
            'sedes': sedes,
            'carreras': carreras,
            'anios': lista_anios,
            'resnro': resnro,
            'total_docs': total_docs,
            'total_docs_baja': total_docs_baja,
            'lista_hs': lista_hs,
            'sedeSel': sedeSel,
            'carreraSel': carreraSel,
            'anioSel': anioSel,
            'fechaDesdeVigSel': fechaDesdeVigSel,
            'fechaHastaVigSel': fechaHastaVigSel,
            'isGet': isGet,
        })

@login_required()
def personasHorasList(request, pid):
    # obtener el Arreglo de años (actual al 2009)
    hoy = date.today()
    anio_actual = hoy.year
    nro_fil = (anio_actual - 2008)
    lista_anios = [None] * nro_fil
    i = 0
    for n in range(nro_fil):
        lista_anios[n] = anio_actual - i
        i += 1
    # Fin obtener el arreglo de años
    form = PersonaHoras()
    perhscat = PersonaHoras.objects.filter(persona_id=pid).order_by("-fecha_inicio")
    persona = Persona.objects.get(id=pid)
    motivosBajaResolucion = ["Fallecimiento", "Renuncia", "Suspensión"]
    if request.method == 'POST':
        id_asignacion = request.POST['id']
        asignacion = PersonaHoras.objects.get(id=id_asignacion)
        fechaBaja = request.POST['fecha_baja']
        if ((str(fechaBaja) >= str(asignacion.fecha_inicio))&(str(fechaBaja)<=str(asignacion.fecha_fin))):
            asignacion.baja = True
            asignacion.resolucion_numero_baja = request.POST['resolucion_numero_baja']
            asignacion.resolucion_anio_baja = request.POST['resolucion_anio_baja']
            asignacion.motivo_baja = request.POST['motivos']
            asignacion.fecha_baja = request.POST['fecha_baja']
            asignacion.observaciones_baja = request.POST['observaciones_baja']
            asignacion.save()
            return render(request, 'personasHorasList.html', {
                'perhscat': perhscat,
                'persona': persona,
                'motivos': motivosBajaResolucion,
                'form': form,
                'anios': lista_anios,
            })
        else:
            errors = "La fecha de baja de la asignación con resolución "+str(asignacion.resolucion_numero)+"/"+str(asignacion.resolucion_anio)+"-UDC debe estar entre "+str(asignacion.fecha_inicio)+" y "+str(asignacion.fecha_fin)
            return render(request, 'personasHorasList.html', {
                'perhscat': perhscat,
                'persona': persona,
                'motivos': motivosBajaResolucion,
                'form': form,
                'anios': lista_anios,
                'error': errors,
            })

    return render(request, 'personasHorasList.html', {
        'perhscat': perhscat,
        'persona': persona,
        'motivos': motivosBajaResolucion,
        'form': form,
        'anios': lista_anios,
    })

@login_required()
def personasHorasBorrarBaja(request, asid):
    asignacion = PersonaHoras.objects.get(id=asid)
    persona = Persona.objects.get(id=asignacion.persona_id)

    asignacion.baja = False
    asignacion.resolucion_numero_baja = 0
    asignacion.resolucion_anio_baja = 0
    asignacion.motivo_baja = ""
    asignacion.fecha_baja = '2000-01-01'
    asignacion.observaciones_baja = ""
    asignacion.save()

    url = '/horascatedras/' + str(persona.id) + '/lista/'
    return redirect(url)

@login_required()
def personasHorasNew(request, pid):
    # obtener el Arreglo de años (actual al 2009)
    hoy = date.today()
    anio_actual = hoy.year
    nro_fil = (anio_actual - 2008)
    lista_anios = [None] * nro_fil
    i = 0
    for n in range(nro_fil):
        lista_anios[n] = anio_actual - i
        i += 1
    # Fin obtener el arreglo de años
    form = PersonaHorasForm()
    persona = Persona.objects.get(id=pid)
    dependencias = Dependencia.objects.all()
    if request.method == 'POST':
        form = PersonaHorasForm(request.POST)
        if form.is_valid():
            perhscat = PersonaHoras.objects.filter(persona_id=pid).order_by("-fecha_inicio")
            persona = Persona.objects.get(id=pid)
            new_perhs = form.save(commit=False)
            new_perhs.persona_id = persona.id
            new_perhs.save()
            return render(request, 'personasHorasList.html', {
                'perhscat': perhscat,
                'persona': persona,
            })
        else:
            form = PersonaHorasForm()
            persona = Persona.objects.get(id=pid)
            return render(request, 'personasHorasNew.html', {
                'persona': persona,
                'form': form,
                'anios': lista_anios,
                'dependencias': dependencias,
            })
    else:
        return render(request, 'personasHorasNew.html', {
            'persona': persona,
            'form': form,
            'anios': lista_anios,
            'dependencias': dependencias,
        })

@login_required()
def personasHorasEdit(request, phid):
    pershoras = PersonaHoras.objects.get(id=phid)
    form = PersonaHorasForm(instance=pershoras)
    persona = Persona.objects.get(id=pershoras.persona_id)
    pid = persona.id
    # obtener el Arreglo de años (actual al 2009)
    hoy = date.today()
    anio_actual = hoy.year
    nro_fil = (anio_actual - 2008)
    lista_anios = [None] * nro_fil
    i = 0
    for n in range(nro_fil):
        lista_anios[n] = anio_actual - i
        i += 1
    # Fin obtener el arreglo de años
    dependencias = Dependencia.objects.all()
    if request.method == 'POST':
        form = PersonaHorasForm(request.POST, instance=pershoras)
        if form.is_valid():
            new_perhs = form.save(commit=False)
            new_perhs.persona_id = persona.id
            new_perhs.save()
            url = '/horascatedras/'+ str(persona.id) + '/lista/'
            return redirect(url)
        else:
            return render(request, 'personasHorasEdit.html', {
                'perhscat': pershoras,
                'persona': persona,
                'form': form,
                'anios': lista_anios,
                'dependencias': dependencias,
            })
    else:
        return render(request, 'personasHorasEdit.html', {
            'perhscat': pershoras,
            'persona': persona,
            'form': form,
            'anios': lista_anios,
            'dependencias': dependencias,
        })

@login_required()
def personasHorasDelete(request, phid):
    pershoras = get_object_or_404(PersonaHoras, id=phid)
    url = '/horascatedras/' + int(phid) + '/lista/'
    if request.method == 'POST':
        pershoras.delete()
        return redirect(url)
    return render(request, url)

@login_required()
def docenteHorasList(request, pid):
    # obtener el Arreglo de años (actual al 2009)
    hoy = date.today()
    anio_actual = hoy.year
    nro_fil = (anio_actual - 2008)
    lista_anios = [None] * nro_fil
    i = 0
    for n in range(nro_fil):
        lista_anios[n] = anio_actual - i
        i += 1
    # Fin obtener el arreglo de años
    form = DocenteHoras()
    motivosBajaResolucion = ["Fallecimiento", "Renuncia", "Suspensión"]
    dochoras = DocenteHoras.objects.filter(persona_id=pid).order_by('-fecha_inicio')
    persona = Persona.objects.get(id=pid)
    nro_fil = len(dochoras)
    lista_hs = [None] * nro_fil
    for n in range(nro_fil):
        lista_hs[n] = {'id_doc': 0, 'tot_hs': 0}
    i = 0
    for dh in dochoras:
        tot_pocentaje = (dh.porcentaje_aplicado * dh.materia.hs_semanales)/100
        tot = math.ceil(dh.materia.hs_semanales + dh.hs_institucionales + tot_pocentaje)
        lista_hs[i]['id_doc'] = dh.id
        lista_hs[i]['tot_hs'] = int(tot)
        i += 1
    if request.method == 'POST':
        id_asignacion = request.POST['id']
        asignacion = DocenteHoras.objects.get(id=id_asignacion)
        fechaBaja = request.POST['fecha_baja']
        if ((str(fechaBaja) >= str(asignacion.fecha_inicio))&(str(fechaBaja)<=str(asignacion.fecha_fin))):
            asignacion.baja = True
            asignacion.resolucion_numero_baja = request.POST['resolucion_numero_baja']
            asignacion.resolucion_anio_baja = request.POST['resolucion_anio_baja']
            asignacion.motivo_baja = request.POST['motivos']
            asignacion.fecha_baja = request.POST['fecha_baja']
            asignacion.observaciones_baja = request.POST['observaciones_baja']
            asignacion.save()
            url = '/horasdocentes/' + str(persona.id) + '/lista/'
            return redirect(url)
        else:
            errors = "La fecha de baja de la asignación con resolución "+str(asignacion.resolucion_numero)+"/"+str(asignacion.resolucion_anio)+"-UDC debe estar entre "+str(asignacion.fecha_inicio)+" y "+str(asignacion.fecha_fin)
            return render(request, 'docentesHorasList.html', {
                'dochoras': dochoras,
                'persona': persona,
                'motivos': motivosBajaResolucion,
                'form': form,
                'anios': lista_anios,
                'lista_hs': lista_hs,
                'error': errors,
            })

    return render(request, 'docentesHorasList.html', {
        'dochoras': dochoras,
        'persona': persona,
        'motivos': motivosBajaResolucion,
        'form': form,
        'anios': lista_anios,
        'lista_hs': lista_hs,
    })

@login_required()
def docenteHorasBorrarBaja(request, asid):
    asignacion = DocenteHoras.objects.get(id=asid)
    persona = Persona.objects.get(id=asignacion.persona_id)

    asignacion.baja = False
    asignacion.resolucion_numero_baja = 0
    asignacion.resolucion_anio_baja = 0
    asignacion.motivo_baja = ""
    asignacion.fecha_baja = '2000-01-01'
    asignacion.observaciones_baja = ""
    asignacion.save()

    url = '/horasdocentes/' + str(persona.id) + '/lista/'
    return redirect(url)

@login_required()
def docenteHorasNew(request, pid):
    lista_remunerado = ['Remunerado', 'Ad Honorem']
    # obtener el Arreglo de años (actual al 2009)
    hoy = date.today()
    anio_actual = hoy.year
    nro_fil = (anio_actual - 2008)
    lista_anios = [None] * nro_fil
    i = 0
    for n in range(nro_fil):
        lista_anios[n] = anio_actual - i
        i += 1
    # Fin obtener el arreglo de años
    sedes = Sede.objects.all()
    carreras = Carrera.objects.all()
    materias = Materia.objects.all().order_by('plan__carrera_id','anio_academico','periodo_id')
    docenteTipos = TipoDocente.objects.all()
    form = DocenteHorasForm()
    persona = Persona.objects.get(id=pid)
    if request.method == 'POST':
        form = DocenteHorasForm(request.POST)
        if form.is_valid():
            mat = Materia.objects.get(id=request.POST['materia'])
            doc_tipo = TipoDocente.objects.get(id=request.POST['docente_tipo'])
            new_dochs = form.save(commit=False)
            new_dochs.persona_id = persona.id
            # new_dochs.fecha_inicio = mat.periodo.fecha_inicio
            # new_dochs.fecha_fin = mat.periodo.fecha_fin
            new_dochs.fecha_inicio = request.POST['fecha_inicio']
            new_dochs.fecha_fin = request.POST['fecha_fin']
            new_dochs.porcentaje_aplicado = doc_tipo.porcentaje_aplicado
            new_dochs.hs_institucionales = doc_tipo.hs_institucionales
            new_dochs.save()
            dochoras = DocenteHoras.objects.filter(persona_id=pid)
            return render(request, 'docentesHorasList.html', {
                'dochoras': dochoras,
                'persona': persona,
                'pid': pid,
            })
        else:
            return render(request, 'docentesHorasNew.html', {
                'persona': persona,
                'form': form,
                'carreras': carreras,
                'sedes': sedes,
                'materias': materias,
                'remunerados': lista_remunerado,
                'docenteTipos': docenteTipos,
                'anios': lista_anios,
            })
    else:
        return render(request, 'docentesHorasNew.html', {
            'persona': persona,
            'form': form,
            'carreras': carreras,
            'sedes': sedes,
            'materias': materias,
            'remunerados': lista_remunerado,
            'docenteTipos': docenteTipos,
            'anios': lista_anios,
        })

@login_required()
def docenteHorasEdit(request, dhid):
    lista_remunerado = ['Remunerado', 'Ad Honorem']
    # obtener el Arreglo de años (actual al 2009)
    hoy = date.today()
    anio_actual = hoy.year
    nro_fil = (anio_actual - 2008)
    lista_anios = [None] * nro_fil
    i = 0
    for n in range(nro_fil):
        lista_anios[n] = str(anio_actual - i)
        i += 1
    # Fin obtener el arreglo de años
    sedes = Sede.objects.all()
    carreras = Carrera.objects.all()
    materias = Materia.objects.all()
    docenteTipos = TipoDocente.objects.all()
    dochoras = DocenteHoras.objects.get(id=dhid)
    form = DocenteHorasForm(instance=dochoras)
    persona = Persona.objects.get(id=dochoras.persona.id)
    if request.method == 'POST':
        form = DocenteHorasForm(request.POST, instance=dochoras)
        if form.is_valid():
            mat = Materia.objects.get(id=request.POST['materia'])
            doc_tipo = TipoDocente.objects.get(id=request.POST['docente_tipo'])
            new_dochs = form.save(commit=False)
            new_dochs.persona_id = persona.id
            #new_dochs.fecha_inicio = mat.periodo.fecha_inicio
            #new_dochs.fecha_fin = mat.periodo.fecha_fin
            new_dochs.fecha_inicio = request.POST['fecha_inicio']
            new_dochs.fecha_fin = request.POST['fecha_fin']
            new_dochs.porcentaje_aplicado = doc_tipo.porcentaje_aplicado
            new_dochs.hs_institucionales = doc_tipo.hs_institucionales
            new_dochs.save()
            dochoras = DocenteHoras.objects.filter(persona_id=persona.id)
            url = '/horasdocentes/'+str(persona.id)+'/lista/'
            return redirect(url)
        else:
            return render(request, 'docentesHorasEdit.html', {
                'dochoras': dochoras,
                'form': form,
                'sedes': sedes,
                'carreras': carreras,
                'docenteTipos': docenteTipos,
                'materias': materias,
                'remunerados': lista_remunerado,
                'anios': lista_anios,
            })
    else:
        return render(request, 'docentesHorasEdit.html', {
            'dochoras': dochoras,
            'form': form,
            'sedes': sedes,
            'carreras': carreras,
            'docenteTipos': docenteTipos,
            'materias': materias,
            'remunerados': lista_remunerado,
            'anios': lista_anios,
        })

@login_required()
def export_admin_xls(request):
    administrativos = gbl_cons_adm
    hoy = date.today()
    dia = hoy.day
    mes = hoy.month
    anio = hoy.year
    str_fecha = str(anio)+'-'+str(mes)+'-'+str(dia)
    response = HttpResponse(content_type='text/xls', charset='utf-8')
    response['Content-Disposition'] = 'attachment; filename="' + str_fecha +'_consulta_horas_catedras.xls"'

    writer = csv.writer(response)
    writer.writerow(['Apellido y Nombre', 'CUIL', 'Resolucion', 'Fecha Inicio', 'Fecha Fin', 'Fecha Baja', 'Hs_catedras', 'Dependencia'])

    for adm in administrativos:
        apenom = (adm.persona.apellidos + ', ' + adm.persona.nombres).encode('ascii', 'replace')
        depend = (adm.dependencia.dependencia_nombre).encode('ascii', 'replace')
        resolu = str(adm.resolucion_numero) + '/' + str(adm.resolucion_anio)
        writer.writerow([apenom, adm.persona.cuil, resolu, adm.fecha_inicio, adm.fecha_fin, adm.fecha_baja, adm.hs_catedras, depend])

    return response

@login_required()
def export_doc_xls(request):
    docentes = gbl_cons_doc
    hoy = date.today()
    dia = hoy.day
    mes = hoy.month
    anio = hoy.year
    str_fecha = str(anio)+'-'+str(mes)+'-'+str(dia)
    response = HttpResponse(content_type='text/xls')
    response['Content-Disposition'] = 'attachment; filename="' + str_fecha +'_consulta_horas_docentes.xls"'

    writer = csv.writer(response)
    writer.writerow(['Apellido y Nombre', 'Sede', 'Carrera', 'Materia', 'ResolNro', 'ResoluAnio', 'Fecha Inicio', 'Fecha Fin',
                     'Fecha Baja', 'Horas Materia', 'Hs total materia', 'Hs a liquidar', 'Anio academico', 'Periodo'])

    for doc in docentes:
        apenom = (doc.persona.apellidos + ', ' + doc.persona.nombres).encode('ascii', 'replace')
        resoluNro = str(doc.resolucion_numero)
        resoluAnio = str(doc.resolucion_anio)
        tot_pocentaje = (doc.porcentaje_aplicado * doc.materia.hs_semanales) / 100
        tot = doc.materia.hs_semanales + doc.hs_institucionales + tot_pocentaje
        writer.writerow(
            [apenom, doc.sede.sede_nombre.encode('ascii', 'replace'), doc.materia.plan.carrera.carrera_nombre.encode('ascii', 'replace')
                , doc.materia.materia_nombre.encode('ascii', 'replace'), resoluNro, resoluAnio, doc.fecha_inicio,
             doc.fecha_fin, doc.fecha_baja, doc.materia.hs_semanales, doc.materia.hs_total_materia, tot,
             doc.materia.anio_academico, doc.materia.periodo.periodo_nombre.encode('ascii', 'replace')]
        )

    return response

@login_required()
def export_cert_prest_serv(request, pid):
    persona = get_object_or_404(Persona, id=pid)
    docHoras = DocenteHoras.objects.filter(persona_id=pid).order_by('fecha_inicio')
    persHoras = PersonaHoras.objects.filter(persona_id=pid).order_by('fecha_inicio')

    lista_dias = {'1': 'primero', '2': 'dos', '3': 'tres', '4': 'cuatro', '5': 'cinco', '6': 'seis', '7': 'siete',
                  '8': 'ocho', '9': 'nueve', '10': 'diez', '11': 'once', '12': 'doce', '13': 'trece',
                  '14': 'catorce', '15': 'quince', '16': 'dieciseis', '17': 'diecisiete', '18': 'dieciocho',
                  '19': 'diecinueve', '20': 'veinte', '21': 'veintiun', '22': 'veintidos', '23': 'veintitres',
                  '24': 'veinticuatro', '25': 'veinticinco', '26': 'veintiseis', '27': 'veintisiete',
                  '28': 'veintiocho', '29': 'veintinueve', '30': 'treinta', '31': 'treinta y uno'}
    lista_mes = {'1': 'enero', '2': 'febrero', '3': 'marzo', '4': 'abril', '5': 'mayo', '6': 'junio',
                 '7': 'julio', '8': 'agosto', '9': 'septiembre', '10': 'octubre', '11': 'noviembre',
                 '12': 'diciembre'}
    lista_anios = {'2008': 'dos mil ocho', '2009': 'dos mil nueve', '2010': 'dos mil diez', '2011': 'dos mil once',
                   '2012': 'dos mil doce', '2013': 'dos mil trece', '2014': 'dos mil catorce', '2015': 'dos mil quince',
                   '2016': 'dos mil dieciseis', '2017': 'dos mil diecisiete', '2018': 'dos mil dieciocho',
                   '2019': 'dos mil diecinueve', '2020': 'dos mil veinte', '2021': 'dos mil vientiuno',
                   '2022': 'dos mil veintidos', '2023': 'dos mil veintitres', '2024': 'dos mil veinticuatro',
                   '2025': 'dos mil veinticinco', '2026': 'dos mil veintiseis', '2027': 'dos mil veintisiete',
                   '2028': 'dos mil veintiocho', '2029': 'dos mil veintinueve', '2030': 'dos mil treinta'}
    hoy = date.today()
    dia = hoy.day
    mes = hoy.month
    anio = hoy.year
    if (dia == 1):
        fecha_en_letras = ' al ' + lista_dias[str(dia)] + ' de ' + lista_mes[str(mes)] + u' del año ' + lista_anios[
            str(anio)]
    else:
        fecha_en_letras = ' a los ' + lista_dias[str(dia)] + u' días del mes de ' + lista_mes[str(mes)] + u' del año ' + \
                          lista_anios[str(anio)]

    docName = "Certificado_Prestacion_de_Servicio_" + persona.apellidos + "-" + persona.nombres + ".docx"

    response = HttpResponse(content_type='text/docx')
    response['Content-Disposition'] = 'attachment; filename="' + docName

    docModelo = finders.find('../static/docs/CertPrestServ_Modelo.docx')
    #file = open(docModelo, 'rb')
    document = Document(docModelo)

    #imagen = finders.find('img/logo_udc.png')
    #document.add_picture(imagen, width=Cm(5), height=Cm(2))
    #renglon = document.paragraphs[-1]
    #renglon.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    #p1 = document.add_paragraph('Rawson, ' + str_fecha)
    #p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p1 = document.add_paragraph(u"Certificado de Prestación de Servicios")

    #document.add_heading("Certificado de Prestación de Servicios", level=1)
    p2 = document.add_paragraph(u'CERTIFICO que el/la Sr./Sra. ')
    p2.add_run(persona.apellidos + ', ' + persona.nombres + ' (C.U.I.L. ' + persona.cuil + ') ').bold = True
    p2.add_run(u'prestó servicios en esta Universidad, desempeñándose en la función y período que a continuación se detallan:')

    table = document.add_table(rows=1, cols=4)
    #table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'Cargo o Función'
    hdr_cells[1].text = u'Resolución'
    hdr_cells[2].text = 'Desde'
    hdr_cells[3].text = 'Hasta'
    if (docHoras.count() > 0):
        for item in docHoras:
            row_cells = table.add_row().cells
            row_cells[0].text = item.docente_tipo.tipo_docente + " de la Asignatura " + \
                                '"' + item.materia.materia_nombre + '"' \
                                + '(' + item.materia.plan.carrera.carrera_nombre + ')'
            row_cells[1].text = u'N° ' + str(item.resolucion_numero) + '/' + str(item.resolucion_anio) + '-UDC'
            #row_cells[2].text = str(item.fecha_inicio).format("d/m/Y")
            #row_cells[3].text = str(item.fecha_fin).format("d/m/Y")
            row_cells[2].text = str(item.fecha_inicio.day)+'/'+str(item.fecha_inicio.month)+'/'+str(item.fecha_inicio.year)
            row_cells[3].text = str(item.fecha_fin.day)+'/'+str(item.fecha_fin.month)+'/'+str(item.fecha_fin.year)

    if (persHoras.count() > 0):
        for item in persHoras:
            row_cells2 = table.add_row().cells
            row_cells2[0].text = item.dependencia.dependencia_nombre
            row_cells2[1].text = u'N° ' + str(item.resolucion_numero) + '/' + str(item.resolucion_anio) + '-UDC'
            #row_cells2[2].text = str(item.fecha_inicio).format('d/m/Y')
            #row_cells2[3].text = str(item.fecha_fin).format('d/m/Y')
            row_cells2[2].text = str(item.fecha_inicio.day)+'/'+str(item.fecha_inicio.month)+'/'+str(item.fecha_inicio.year)
            row_cells2[3].text = str(item.fecha_fin.day)+'/'+str(item.fecha_fin.month)+'/'+str(item.fecha_fin.year)
    p20 = document.add_paragraph('')
    p3 = document.add_paragraph('Licencia SIN GOCE DE HABERES: ')
    p3.add_run('NO REGISTRA.').font.color.rgb = RGBColor(248, 000, 000)

    p4 = document.add_paragraph('Aportes en el Instituto de Seguridad Social y Seguros - Chubut.')

    p5 = document.add_paragraph('A pedido de el/la interesado/a y a solo efectos de acreditar haber prestado'
                                + u' servicios en nuestra institución, se extiende el presente en Rawson (Chubut)'
                                + ' a los' + fecha_en_letras)
    document.save(response)

    return response
